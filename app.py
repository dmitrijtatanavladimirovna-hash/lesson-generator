import streamlit as st
import google.generativeai as genai
import io
from docx import Document

# Настройка страницы
st.set_page_config(page_title="Конструктор КСП", layout="wide")
st.title("🎓 Конструктор КСП (Краткосрочный план)")

# 1. Настройка API
if "GOOGLE_API_KEY" not in st.secrets:
    st.error("❌ Ключ GOOGLE_API_KEY не найден. Перейди в Settings -> Secrets.")
    st.stop()

genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

try:
    model = genai.GenerativeModel('gemini-1.5-flash')
except Exception as e:
    st.error(f"Ошибка настройки модели: {e}")
    st.stop()

# 2. Боковая панель (Данные педагога)
with st.sidebar:
    st.header("Данные педагога")
    fio = st.text_input("ФИО педагога")
    school = st.text_input("Школа")
    subject = st.selectbox("Предмет", ["Математика", "Алгебра", "Геометрия", "Физика", "Химия", "Биология", "История", "География", "Информатика", "Другое"])
    grade = st.selectbox("Класс", [f"{i} класс" for i in range(1, 12)])

# 3. Ввод данных
tab1, tab2 = st.tabs(["📋 Ввод данных", "📄 Результат"])

with tab1:
    topic = st.text_input("Тема урока")
    col1, col2 = st.columns(2)
    with col1:
        objectives = st.text_area("Цели обучения (из КТП)")
        differentiation = st.text_input("Дифференциация")
    with col2:
        evaluation = st.text_area("Критерии оценивания")
        reflection = st.text_area("Рефлексия")

    if st.button("Сгенерировать план"):
        if not topic or not objectives:
            st.warning("Заполните тему и цели!")
        else:
            with st.spinner('Пишу план...'):
                prompt = f"""
                Составь подробный КСП (Краткосрочный план) по ГОСО.
                Учитель: {fio}. Школа: {school}. Предмет: {subject}. Класс: {grade}.
                Тема: {topic}.
                Цели обучения: {objectives}.
                Дифференциация: {differentiation}.
                Критерии оценивания: {evaluation}.
                Рефлексия: {reflection}.

                СТРУКТУРА:
                1. Цели урока (SMART).
                2. Этапы урока (Таблица: Время, Действия, Оценивание).
                3. Ресурсы.
                4. Дифференциация.
                5. Рефлексия.
                """
                try:
                    response = model.generate_content(prompt)
                    st.session_state['plan'] = response.text
                    st.success("План готов! Перейди во вкладку 'Результат'.")
                except Exception as e:
                    st.error(f"Ошибка при генерации: {e}")

# 4. Результат и Word
with tab2:
    if 'plan' in st.session_state:
        st.markdown(st.session_state['plan'])
        
        # Генерация Word
        doc = Document()
        doc.add_heading(f'КСП: {topic}', 0)
        doc.add_paragraph(st.session_state['plan'])
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label="📥 Скачать в Word",
            data=bio.getvalue(),
            file_name=f"{topic}_КСП.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Сначала заполни данные во вкладке 'Ввод данных'")
